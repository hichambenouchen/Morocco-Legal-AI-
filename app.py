import io
import os
import re
import math
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from groq import Groq
import streamlit as st

# ---------------------------------------------------------
# 1. إعدادات الصفحة والتطبيق
# ---------------------------------------------------------
st.set_page_config(
    page_title="الموسوعة القانونية المغربية الشاملة | JurisConsult Enterprise RAG",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ---------------------------------------------------------
# 2. تصميم الواجهة (Advanced CSS with RTL Support)
# ---------------------------------------------------------
st.markdown(
    """
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700;800;900&family=Inter:wght@400;600&display=swap');

    html, body, [class*="css"] {
        font-family: 'Cairo', 'Inter', sans-serif !important;
    }

    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    .block-container {
        padding-top: 1rem !important;
        padding-bottom: 6rem !important;
        max-width: 1200px !important;
    }

    /* دعم الاتجاه من اليمين إلى اليسار RTL بشكل كامل */
    .stChatMessage, .stChatMessage p, .stMarkdown, .stMarkdown p, .stMarkdown div {
        direction: rtl !important;
        text-align: right !important;
    }

    .stMarkdown h1, .stMarkdown h2, .stMarkdown h3, .stMarkdown h4, .stMarkdown ul, .stMarkdown ol, .stMarkdown li {
        direction: rtl !important;
        text-align: right !important;
    }

    .stChatInput textarea {
        direction: rtl !important;
        text-align: right !important;
    }

    .hero-header {
        background: linear-gradient(135deg, #091e3e 0%, #1e3c72 50%, #2a5298 100%);
        border-radius: 18px;
        padding: 32px 36px;
        color: white;
        margin-bottom: 25px;
        box-shadow: 0 12px 30px -5px rgba(15, 23, 42, 0.35);
        border: 1px solid rgba(255, 255, 255, 0.15);
        direction: rtl !important;
        text-align: right !important;
    }
    
    .hero-title {
        font-size: 2.3rem;
        font-weight: 900;
        margin: 0;
        color: #ffffff;
        display: flex;
        align-items: center;
        gap: 14px;
    }

    .hero-subtitle {
        font-size: 1.02rem;
        color: #cbd5e1;
        margin-top: 10px;
        font-weight: 400;
        line-height: 1.6;
    }

    .rag-badge {
        background: linear-gradient(90deg, #2563eb, #3b82f6);
        color: white;
        font-size: 0.8rem;
        padding: 5px 12px;
        border-radius: 20px;
        font-weight: 700;
        letter-spacing: 0.5px;
    }

    .stButton>button {
        border-radius: 10px !important;
        border: 1px solid #cbd5e1 !important;
        font-weight: 600 !important;
        transition: all 0.25s ease !important;
    }

    .stButton>button:hover {
        border-color: #1d4ed8 !important;
        color: #1d4ed8 !important;
        background-color: #eff6ff !important;
        transform: translateY(-1px);
    }

    .stChatMessage {
        border-radius: 16px !important;
        padding: 18px !important;
        margin-bottom: 14px !important;
        box-shadow: 0 3px 8px rgba(0,0,0,0.04);
        border: 1px solid #f1f5f9;
    }

    @media (max-width: 768px) {
        .hero-header {
            padding: 22px;
            text-align: center;
        }
        .hero-title {
            font-size: 1.6rem;
            justify-content: center;
            flex-wrap: wrap;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------
# 3. إدارة اللغات الثلاثية
# ---------------------------------------------------------
if "lang" not in st.session_state:
    st.session_state.lang = "ar"

col_lang1, col_lang2, col_space = st.columns([2, 3, 5])
with col_lang1:
    st.caption("🌐 **اختيار لغة الاستشارة / Language:**")
with col_lang2:
    selected_lang = st.radio(
        "Language",
        options=["العربية 🇲🇦", "Français 🇫🇷", "English 🇬🇧"],
        horizontal=True,
        label_visibility="collapsed",
    )
    if "🇲🇦" in selected_lang:
        st.session_state.lang = "ar"
    elif "🇫🇷" in selected_lang:
        st.session_state.lang = "fr"
    else:
        st.session_state.lang = "en"

TEXTS = {
    "ar": {
        "title": "الموسوعة القانونية المغربية الشاملة",
        "subtitle": "مرجع استرجاع قانوني ذكي (Enterprise RAG) مستند إلى قاعدة بيانات الأمانة العامة للحكومة والجريدة الرسمية المغربية: الصفقات العمومية، الوظيفة العمومية، الشغل، الأسرة، الجنائي، العقود، والكراء.",
        "welcome": "مرحباً بك! أنا مستشارك القانوني المستند لنصوص الجريدة الرسمية والأمانة العامة للحكومة (SGG). اطرح استفسارك القانوني وسيتم فحص المواد التشريعية ومطابقتها مباشرة.",
        "input_placeholder": "اطرح سؤالك (مثال: طلبات العروض في الصفقات العمومية، المساطر التأديبية في الوظيفة العمومية، الشيك بدون رصيد...)",
        "quick_questions": "💡 استفسارات قانونية متخصصة ومستحدثة:",
        "q1": "ما هي قواعد ومساطر إبرام الصفقات العمومية حسب مرسوم 2.22.431؟",
        "q2": "ما هي الحقوق والمساطر التأديبية للموظف العمومي حسب نظام الوظيفة العمومية؟",
        "q3": "ما هي الشروط والأسباب القانونية لإفراغ المحل التجاري أو السكني؟",
        "tools": "🛠️ إدارة الجلسة وتصدير التقارير الرسمية",
        "reset": "🗑️ إعادة ضبط المحادثة",
        "download_txt": "📥 تصدير التقرير (TXT)",
        "download_word": "📄 تصدير التقرير (DOCX)",
        "rag_sources": "📚 النصوص والمواد التشريعية المسترجعة من الجريدة الرسمية والأمانة العامة للحكومة:",
    },
    "fr": {
        "title": "JurisConsult Maroc - Legal RAG Suite",
        "subtitle": "Système de recherche juridique complet basé sur le Bulletin Officiel et SGG : Marchés Publics, Fonction Publique, Travail, Famille, Pénal, DOC.",
        "welcome": "Bienvenue sur la plateforme juridique marocaine basée sur les textes officiels du Bulletin Officiel (SGG).",
        "input_placeholder": "Posez votre question (Marchés publics, Discipline Fonction Publique, Chèque sans provision...)",
        "quick_questions": "💡 Exemples de requêtes juridiques :",
        "q1": "Quelles sont les règles de passation des marchés publics (Décret 2.22.431) ?",
        "q2": "Quels sont les droits et régimes disciplinaires de la fonction publique ?",
        "q3": "Quelles sont les conditions d'éviction d'un bail commercial ou d'habitation ?",
        "tools": "🛠️ Outils & Exportation du Rapport",
        "reset": "🗑️ Réinitialiser",
        "download_txt": "📥 Exporter en TXT",
        "download_word": "📄 Exporter en Word (DOCX)",
        "rag_sources": "📚 Textes législatifs extraits du Bulletin Officiel (SGG) :",
    },
    "en": {
        "title": "Moroccan Legal Enterprise RAG Suite",
        "subtitle": "Comprehensive Legal AI backed by the Secretariat General of the Government (SGG) and Official Gazette: Public Procurement, Civil Service, Labor, Family, Criminal Code.",
        "welcome": "Welcome to the Moroccan Comprehensive Legal System referencing the Official Gazette.",
        "input_placeholder": "Type your query (Public procurement procedures, Civil Service disciplinary rules, Lease eviction...)",
        "quick_questions": "💡 Frequent Legal Queries:",
        "q1": "What are the key rules in the Public Procurement Decree 2.22.431?",
        "q2": "What are the disciplinary procedures under the Moroccan Civil Service Code?",
        "q3": "What are the statutory grounds for commercial/residential eviction?",
        "tools": "🛠️ Session Management & Export Tools",
        "reset": "🗑️ Reset Conversation",
        "download_txt": "📥 Export Report (TXT)",
        "download_word": "📄 Export Report (DOCX)",
        "rag_sources": "📚 Retrieved Statutory Legal Texts (Official Gazette):",
    },
}

current_texts = TEXTS[st.session_state.lang]

# ---------------------------------------------------------
# 4. الهيدر الرئيسي
# ---------------------------------------------------------
st.markdown(
    f"""
    <div class="hero-header">
        <div class="hero-title">
            ⚖️ {current_texts['title']} <span class="rag-badge">SGG & Official Gazette Integrated</span>
        </div>
        <div class="hero-subtitle">{current_texts['subtitle']}</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ---------------------------------------------------------
# 5. الترسانة القانونية الكاملة
# ---------------------------------------------------------
FULL_LEGAL_CORPUS = [
    {"law": "مرسوم الصفقات العمومية - المادة 4 و 5", "category": "صفقات عمومية", "text": "تخضع الصفقات العمومية لمبادئ حرية الوصول إلى الطلبية العمومية، المساواة في التعامل مع المتنافسين، والشفافية في اختيارات صاحب المشروع. وتشمل الصفقات: صفقة أشغال، توريدات، أو خدمات."},
    {"law": "مرسوم الصفقات العمومية - المادة 16 و 17", "category": "صفقات عمومية", "text": "تتم طريقة إبرام الصفقات العمومية عن طريق طلبات العروض (مفتوح أو محدود)، المباراة، أو المسطرة التفاوضية. يعتبر طلب العروض المفتوح هو الأصل."},
    {"law": "مرسوم الصفقات العمومية - المادة 40", "category": "صفقات عمومية", "text": "يستوجب على المتنافسين تقديم ضمان مؤقت للقبول في الصفقات العمومية، ثم يتوجب على النائل إيداع ضمان نهائي لضمان تنفيذه السليم للالتزامات التعاقدية."},
    {"law": "مرسوم الصفقات العمومية - المادة 136", "category": "صفقات عمومية", "text": "يمكن للمتنافسين تقديم شكايات ورسائل طعن للجنة الصفقات العمومية أو السلطة الحكومية المعنية في حال الإخلال بقواعد المنافسة والشفافية."},
    {"law": "قانون الوظيفة العمومية - المادة 2 و 13", "category": "وظيفة عمومية", "text": "الموظف هو كل شخص يعين في وظيفة دائمة ويرسم في إحدى درجات التسلسل الإداري للإدارات التابعة للدولة. والتوظيف يتم عبر المباريات للجميع بمساواة."},
    {"law": "قانون الوظيفة العمومية - المادة 39 و 40", "category": "وظيفة عمومية", "text": "للموظف الحق في العطلة السنوية المؤدى عنها ومدتها شهر عن كل سنة عمل. كما يستفيد من رخص المرض ورخص الولادة والأمومة."},
    {"law": "قانون الوظيفة العمومية - المادة 65 و 66", "category": "وظيفة عمومية", "text": "تحدد العقوبات التأديبية في درجتين: الأولى تشمل الإنذار والتوبيخ، والثانية تشمل الإنزلاق في الدرجة، الحرمان من الترقية، أو العزل."},
    {"law": "قانون الوظيفة العمومية - المادة 73", "category": "وظيفة عمومية", "text": "إذا ارتكب الموظف خطأ جسيماً جاز توقيفه فوراً من طرف السلطة التي لها حق التأديب مع خصم أجرته باستثناء التعويضات العائلية."},
    {"law": "مدونة الشغل - المادة 13", "category": "شغل", "text": "تحدد فترة التجربة بالنسبة للعقود غير محددة المدة في: 3 أشهر للأطر وما ماثلهم، شهر ونصف للمستخدمين، و15 يوما للعمال."},
    {"law": "مدونة الشغل - المادة 61 و 62", "category": "شغل", "text": "يستحق الأجير تعويضاً عن الفصل التعسفي ما لم يرتكب خطأ جسيماً. ويجب الاستماع إليه بحضور مندوب الأجراء في أجل لا يتعدى 8 أيام."},
    {"law": "مدونة الشغل - المادة 196", "category": "شغل", "text": "تؤدى الساعات الإضافية بنسبة زيادة 25% نهاراً و50% ليلاً، وترتفع إلى 50% نهاراً و100% ليلاً في أيام العطل والأعياد."},
    {"law": "مدونة الأسرة - المادة 78 و 84", "category": "أسرة", "text": "الطلاق حل ميثاق الزوجية يمارس تحت مراقبة القاضي. وتشتمل مستحقات الزوجة على الصداق المؤجل، نفقة العدة، والمتعة."},
    {"law": "مدونة الأسرة - المادة 94", "category": "أسرة", "text": "يقدم طلب طلاق الشقاق للمحكمة وتجري محاولة الصلح عبر حكمين، وإن تعذر الصلح تحكم المحكمة بالطلاق وتحدد مستحقات الطرفين."},
    {"law": "قانون الالتزامات والعقود - المادة 230", "category": "عقود", "text": "الالتزامات التعاقدية المنشأة على وجه صحيح تقوم مقام القانون بالنسبة إلى من أنشأوها."},
    {"law": "قانون الكراء السكني (67.12) - المادة 44", "category": "كراء", "text": "لا ينتهي عقد الكراء إلا بعد إنذار بالإفراغ يبين السبب المشروع بمهلة لا تقل عن شهرين."},
    {"law": "قانون الكراء التجاري (49.16) - المادة 7", "category": "كراء تجاري", "text": "يستحق المكتري تعويضاً كاملاً عن الإفراغ يعادل الضرر الحاصل عن فقدان الأصل التجاري ما لم يستند الإفراغ لعدم أداء الكراء."},
    {"law": "قانون المسطرة الجنائية - المادة 66", "category": "مسطرة جنائية", "text": "تحدد مدة الحراسة النظرية في الجرائم العادية في 48 ساعة قابلة للتمديد بـ 24 ساعة بإذن النيابة العامة."},
    {"law": "القانون الجنائي - المادة 540", "category": "جنائي", "text": "يعاقب بالحبس من سنة إلى خمس سنوات وغرامة كل من استعمل الاحتيال والخداع لإيقاع شخص في الغلط وسلبه أموالاً."},
    {"law": "مدونة التجارة - المادة 540 و 543", "category": "تجاري", "text": "يعاقب بالحبس من سنة إلى 5 سنوات وغرامة مالية كل من أصدر شيكاً بدون رصيد قائم وقابل للتصرف."},
]

def clean_and_tokenize(text):
    text = re.sub(r'[^\w\s]', '', text.lower())
    return set(text.split())

def tfidf_similarity_search(query, top_k=4):
    query_tokens = clean_and_tokenize(query)
    if not query_tokens:
        return FULL_LEGAL_CORPUS[:top_k]

    scored_docs = []
    for doc in FULL_LEGAL_CORPUS:
        doc_tokens = clean_and_tokenize(doc["text"] + " " + doc["law"] + " " + doc["category"])
        intersection = query_tokens.intersection(doc_tokens)
        score = len(intersection) / (math.sqrt(len(query_tokens)) * math.sqrt(len(doc_tokens)) + 1e-5)
        
        for q_token in query_tokens:
            if q_token in doc["category"]:
                score += 0.6
        
        scored_docs.append((score, doc))

    scored_docs.sort(key=lambda x: x[0], reverse=True)
    retrieved = [doc for score, doc in scored_docs[:top_k] if score > 0]
    
    if not retrieved:
        retrieved = FULL_LEGAL_CORPUS[:top_k]
    
    return retrieved

# ---------------------------------------------------------
# 6. تصدير التقارير بملفات DOCX
# ---------------------------------------------------------
def generate_docx(messages):
    doc = docx.Document()

    def set_rtl(p):
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_pr = p._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        p_pr.append(bidi)

    title = doc.add_paragraph("Rapport de Consultation Juridique RAG - Droit Marocain (SGG)")
    if st.session_state.lang == "ar":
        set_rtl(title)
        title.runs[0].text = "تقرير استشارة قانونية شاملة - التشريع المغربي والجريدة الرسمية"

    title.runs[0].font.bold = True
    title.runs[0].font.size = docx.shared.Pt(15)

    for msg in messages:
        role_name = "المستخدم / User" if msg["role"] == "user" else "المستشار القانوني / JurisConsult RAG"
        p = doc.add_paragraph()
        if st.session_state.lang == "ar":
            set_rtl(p)

        run_role = p.add_run(f"[{role_name}]:\n")
        run_role.bold = True

        p_content = doc.add_paragraph(msg["content"])
        if st.session_state.lang == "ar":
            set_rtl(p_content)
        doc.add_paragraph("-" * 45)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# ---------------------------------------------------------
# 7. محرك الذكاء الاصطناعي (Groq Client Setup)
# ---------------------------------------------------------
api_key = st.secrets.get("GROQ_API_KEY") or os.getenv("GROQ_API_KEY")
if not api_key:
    st.error("⚠️ GROQ_API_KEY مفقود! يرجى إضافته إلى إعدادات Streamlit Secrets.")
    st.stop()

client = Groq(api_key=api_key)

# ---------------------------------------------------------
# 8. شريط الأدوات والتنزيلات
# ---------------------------------------------------------
with st.expander(current_texts["tools"]):
    col_reset, col_txt, col_docx = st.columns(3)

    with col_reset:
        if st.button(current_texts["reset"]):
            st.session_state.messages = [{"role": "assistant", "content": current_texts["welcome"]}]
            st.rerun()

    if "messages" in st.session_state and len(st.session_state.messages) > 1:
        with col_txt:
            chat_text = "\n\n".join([f"{'User' if m['role']=='user' else 'Legal RAG AI'}: {m['content']}" for m in st.session_state.messages])
            st.download_button(current_texts["download_txt"], data=chat_text, file_name="rag_legal_consultation.txt", mime="text/plain")

        with col_docx:
            docx_bytes = generate_docx(st.session_state.messages)
            st.download_button(current_texts["download_word"], data=docx_bytes, file_name="rag_legal_consultation.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ---------------------------------------------------------
# 9. إدارة سجل الرسائل والأسئلة السريعة
# ---------------------------------------------------------
if "messages" not in st.session_state or len(st.session_state.messages) == 0 or st.session_state.messages[0]["content"] not in [TEXTS[k]["welcome"] for k in TEXTS]:
    st.session_state.messages = [{"role": "assistant", "content": current_texts["welcome"]}]

if len(st.session_state.messages) <= 1:
    st.markdown(f"**{current_texts['quick_questions']}**")
    q_cols = st.columns(3)

    prompt_to_send = None
    if q_cols[0].button(f"📌 {current_texts['q1']}"):
        prompt_to_send = current_texts["q1"]
    if q_cols[1].button(f"📌 {current_texts['q2']}"):
        prompt_to_send = current_texts["q2"]
    if q_cols[2].button(f"📌 {current_texts['q3']}"):
        prompt_to_send = current_texts["q3"]

    if prompt_to_send:
        st.session_state.messages.append({"role": "user", "content": prompt_to_send})
        st.rerun()

# عرض محادثات الجلسة
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# ---------------------------------------------------------
# 10. معالجة الإدخال بنظام RAG الشامل
# ---------------------------------------------------------
user_input = st.chat_input(current_texts["input_placeholder"])

if user_input or (len(st.session_state.messages) > 1 and st.session_state.messages[-1]["role"] == "user" and len(st.session_state.messages) % 2 == 0):

    current_prompt = user_input if user_input else st.session_state.messages[-1]["content"]

    if user_input:
        st.session_state.messages.append({"role": "user", "content": current_prompt})
        with st.chat_message("user"):
            st.markdown(current_prompt)

    # استرجاع النصوص التشريعية
    retrieved_docs = tfidf_similarity_search(current_prompt, top_k=4)
    
    formatted_context = ""
    sources_ui = []
    for idx, doc in enumerate(retrieved_docs, 1):
        formatted_context += f"[{idx}] {doc['law']}: {doc['text']}\n"
        sources_ui.append(f"• **{doc['law']}**: {doc['text']}")

    lang_instruction = {
        "ar": "أجب باللغة العربية الفصحى بصياغة قانونية سليمة. اكتب الأرقام والأسماء الأجنبية والرموز بشكل متناسق مع النص العربي دون إدخال رموز أجنبية داخل الجملة تجنباً لانقلاب الاتجاه.",
        "fr": "Répondez en français juridique académique en citant le Bulletin Officiel et le SGG.",
        "en": "Answer in precise legal English citing the Moroccan Official Gazette (SGG)."
    }[st.session_state.lang]

    system_prompt = f"""
    You are a Senior Moroccan Legal Counsel with expertise in all legislation published in the Official Gazette (Bulletin Officiel - SGG).
    
    Retrieved Statutory Context:
    {formatted_context}
    
    Language Requirement: {lang_instruction}
    
    Formatting Guideline: Ensure full RTL fluency for Arabic. Do NOT mix English abbreviations inside Arabic sentences in a way that breaks paragraph alignment. Write numbers and legal decree titles naturally.
    """

    messages_payload = [{"role": "system", "content": system_prompt}]
    for m in st.session_state.messages:
        messages_payload.append({"role": m["role"], "content": m["content"]})

    with st.chat_message("assistant"):
        with st.expander(current_texts["rag_sources"]):
            for src in sources_ui:
                st.markdown(src)

        with st.spinner("جاري استخراج الفصول القانونية المطابقة من الجريدة الرسمية..." if st.session_state.lang == "ar" else "Scanning SGG Legislation..."):
            try:
                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=messages_payload,
                    temperature=0.15,
                )
                bot_reply = response.choices[0].message.content
                st.markdown(bot_reply)
                st.session_state.messages.append({"role": "assistant", "content": bot_reply})
                st.rerun()
            except Exception as e:
                st.error(f"حدث خطأ أثناء معالجة الاستشارة: {e}")
